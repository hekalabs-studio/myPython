from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def add_code_block(doc, code_text):
    """Menambahkan blok kode dengan latar belakang abu-abu dan font monospace"""
    table = doc.add_table(rows=1, cols=1)
    # Atur warna latar belakang abu-abu muda (F3F4F6)
    shading_elm = parse_xml(r'<w:shd {} w:fill="F3F4F6"/>'.format(nsdecls('w')))
    table.cell(0, 0)._tc.get_or_add_tcPr().append(shading_elm)
    
    p = table.cell(0, 0).paragraphs[0]
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)

def add_callout(doc, text, callout_type="tip"):
    """Menambahkan kotak Tips atau Peringatan dengan warna latar"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    
    # Warna latar & teks berdasarkan tipe
    if callout_type == "tip":
        bg_color = "F0F7FF" # Biru sangat muda
        title_text = "💡 TIPS ROBOTIKA: "
        title_color = RGBColor(0, 102, 204)
    else:
        bg_color = "FFF4E5" # Oranye sangat muda
        title_text = "⚠️ PENTING: "
        title_color = RGBColor(204, 85, 0)
        
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
    p.paragraph_format.element.get_or_add_pPr().append(shading_elm)
    
    run_title = p.add_run(title_text)
    run_title.bold = True
    run_title.font.color.rgb = title_color
    
    run_text = p.add_run(text)
    run_text.font.color.rgb = RGBColor(50, 50, 50)

def add_styled_heading(doc, text, level, color_rgb):
    """Menambahkan heading dengan warna kustom"""
    if level == 1:
        p = doc.add_heading(text, level=1)
    elif level == 2:
        p = doc.add_heading(text, level=2)
    else:
        p = doc.add_heading(text, level=3)
        
    for run in p.runs:
        run.font.color.rgb = color_rgb

def add_styled_table(doc, headers, data):
    """Menambahkan tabel dengan style yang rapi"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading Accent 1'
    
    # Header
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                
    # Data
    for row_data in data:
        row_cells = table.add_row().cells
        for i, item in enumerate(row_data):
            row_cells[i].text = str(item)

def main():
    doc = Document()
    
    # --- HALAMAN JUDUL ---
    title = doc.add_heading('MODUL BELAJAR: C++ dari Nol hingga Robotika', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
        run.font.size = Pt(24)
        
    subtitle = doc.add_paragraph('Dirancang Khusus untuk Mahasiswa Teknik Elektro & Robotika\nEdisi Lengkap, Diperbarui, dan Visual | 2026')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_page_break()

    # --- DAFTAR ISI ---
    add_styled_heading(doc, '📑 DAFTAR ISI', 1, RGBColor(0, 51, 102))
    toc_data = [
        ["🟢 Level 1", "Fondasi C++", "Variabel, I/O, Kondisional, Perulangan, Fungsi, Array"],
        ["🟡 Level 2", "Struktur Data & STL", "Vector, Map, Algorithm, String"],
        ["🟠 Level 3", "Pointer & Memori", "Pointer, Reference, Dynamic Memory, Smart Pointer"],
        ["🔴 Level 4", "OOP & File Handling", "Class, Inheritance, Polymorphism, File I/O"],
        ["🟣 Level 5", "Robotika & Embedded", "Bit Manipulation, Multithreading, Serial Comm, Proyek"],
        ["🛠️ Lampiran", "Checklist & Tools", "Checklist Belajar & Rekomendasi Software"]
    ]
    add_styled_table(doc, ["Level", "Topik", "Keterangan"], toc_data)
    doc.add_paragraph("\n(Catatan: Di Microsoft Word, kamu bisa memperbarui ini menjadi Daftar Isi Otomatis via Tab 'References' > 'Table of Contents')\n")
    doc.add_page_break()

    # ==========================================
    # LEVEL 1
    # ==========================================
    add_styled_heading(doc, '🟢 LEVEL 1 — FONDASI C++', 1, RGBColor(34, 139, 34))
    doc.add_paragraph('Level ini adalah pondasi dari segalanya. Jangan dilewati! Semua kode robotika dibangun di atas konsep-konsep berikut.')

    add_styled_heading(doc, '1.1 Variabel & Tipe Data', 2, RGBColor(34, 139, 34))
    doc.add_paragraph('Variabel adalah "wadah" untuk menyimpan data. Setiap variabel memiliki tipe yang menentukan jenis data dan ukuran memori.')
    
    code_1_1 = """#include <iostream>
#include <string>
using namespace std;

int main() {
    int umur = 20;                  
    long long jarak = 9000000000LL; 
    unsigned int pin = 13;          

    float suhu = 36.5f;             
    double tegangan = 3.3;          

    char grade = 'A';               
    string nama = "Budi";           

    bool motorHidup = true;

    cout << "Nama : " << nama << endl;
    cout << "Suhu : " << suhu << " C" << endl;
    return 0;
}"""
    add_code_block(doc, code_1_1)
    
    add_callout(doc, "Gunakan 'int' untuk nomor pin, 'float'/'double' untuk nilai sensor, dan 'bool' untuk status relay/motor.", "tip")
    add_callout(doc, "Selalu tambahkan 'f' di belakang nilai float (3.14f). Jika tidak, compiler akan menganggapnya double dan membuang memori.", "warning")

    table_tipe_data = [
        ["short", "2 byte", "-32.768 s/d 32.767", "short suhu = -10;"],
        ["int", "4 byte", "-2 M s/d 2 M", "int pin = 13;"],
        ["long long", "8 byte", "Sangat besar", "long long data = 9e18;"],
        ["unsigned int", "4 byte", "0 s/d 4 M", "unsigned int n = 300;"],
        ["float", "4 byte", "6-7 digit presisi", "float pi = 3.14f;"],
        ["double", "8 byte", "15-16 digit presisi", "double g = 9.81;"],
        ["char", "1 byte", "1 karakter ASCII", "char c = 'A';"],
        ["string", "Dinamis", "Teks bebas", "string s = \"hello\";"],
        ["bool", "1 byte", "true / false", "bool nyala = true;"]
    ]
    add_styled_table(doc, ["Tipe Data", "Ukuran", "Rentang Nilai", "Contoh Penggunaan Robotika"], table_tipe_data)

    # ==========================================
    # LEVEL 2
    # ==========================================
    add_styled_heading(doc, '🟡 LEVEL 2 — STRUKTUR DATA & STL', 1, RGBColor(218, 165, 32))
    doc.add_paragraph('STL (Standard Template Library) adalah "senjata rahasia" C++. Lebih aman dan fleksibel daripada array biasa.')

    add_styled_heading(doc, '2.1 Vector (Array Dinamis)', 2, RGBColor(218, 165, 32))
    doc.add_paragraph('Ukurannya bisa berubah saat program berjalan. Sangat cocok untuk menyimpan data sensor yang jumlahnya tidak pasti.')
    
    code_2_1 = """#include <vector>
#include <iostream>
using namespace std;

int main() {
    vector<float> dataSensor; 
    
    dataSensor.push_back(3.14f);
    dataSensor.push_back(2.71f);
    
    cout << "Data pertama: " << dataSensor.front() << endl;
    cout << "Jumlah data: " << dataSensor.size() << endl;
    
    dataSensor.pop_back(); // Hapus data terakhir
    return 0;
}"""
    add_code_block(doc, code_2_1)

    add_styled_heading(doc, '2.2 Library <algorithm>', 2, RGBColor(218, 165, 32))
    doc.add_paragraph('Fungsi siap pakai untuk mengolah data, seperti sorting dan searching.')
    
    code_2_2 = """#include <algorithm>
#include <vector>
using namespace std;

int main() {
    vector<int> data = {5, 2, 8, 1, 9, 3};
    
    // 1. Sorting (Mengurutkan)
    sort(data.begin(), data.end()); // Ascending
    
    // 2. Mencari elemen
    auto it = find(data.begin(), data.end(), 8);
    if (it != data.end()) {
        cout << "Ditemukan di index: " << distance(data.begin(), it) << endl;
    }
    return 0;
}"""
    add_code_block(doc, code_2_2)

    # ==========================================
    # LEVEL 3
    # ==========================================
    add_styled_heading(doc, '🟠 LEVEL 3 — POINTER & MANAJEMEN MEMORI', 1, RGBColor(255, 140, 0))
    doc.add_paragraph('Fitur paling kuat di C++. Memungkinkan akses langsung ke alamat memori hardware, yang merupakan inti dari Embedded System.')

    add_styled_heading(doc, '3.1 Pointer & Reference', 2, RGBColor(255, 140, 0))
    
    code_3_1 = """int main() {
    int x = 42;
    int* ptr = &x;  // '&' = ambil alamat dari x
    
    cout << "Nilai x    : " << x << endl;    // 42
    cout << "Alamat x   : " << &x << endl;   // 0x7fff...
    cout << "Isi ptr    : " << ptr << endl;  // 0x7fff...
    cout << "Nilai di ptr: " << *ptr << endl;// 42 ('*' = dereference)
    
    *ptr = 100; // Mengubah nilai x MELALUI pointer
    cout << "Nilai x sekarang: " << x << endl; // 100
    return 0;
}"""
    add_code_block(doc, code_3_1)
    
    add_callout(doc, "Gunakan Reference (int& x) daripada Pointer untuk passing parameter ke fungsi jika memungkinkan, karena lebih aman dari null pointer.", "tip")

    add_styled_heading(doc, '3.2 Smart Pointer (C++11)', 2, RGBColor(255, 140, 0))
    doc.add_paragraph('Solusi modern untuk mencegah Memory Leak. Memori dihapus otomatis.')
    
    code_3_2 = """#include <memory>
#include <iostream>
using namespace std;

int main() {
    // unique_ptr: Otomatis dihapus saat keluar scope
    {
        unique_ptr<int> ptr = make_unique<int>(42);
        cout << "Nilai: " << *ptr << endl;
    } // <-- Di sini, memori otomatis dihapus, tidak perlu 'delete' manual!
    return 0;
}"""
    add_code_block(doc, code_3_2)

    # ==========================================
    # LEVEL 4
    # ==========================================
    add_styled_heading(doc, '🔴 LEVEL 4 — OOP (OBJECT-ORIENTED PROGRAMMING)', 1, RGBColor(220, 20, 60))
    doc.add_paragraph('Membuat kode terstruktur seperti dunia nyata. Robot bukan sekadar kumpulan variabel, tapi objek yang punya properti dan perilaku.')

    add_styled_heading(doc, '4.1 Class & Object', 2, RGBColor(220, 20, 60))
    
    code_4_1 = """class MotorDC {
private:
    int pin;
    int kecepatan;
public:
    MotorDC(int p) { // Constructor
        pin = p;
        kecepatan = 0;
    }
    
    ~MotorDC() { // Destructor
        cout << "Motor dimatikan." << endl;
    }
    
    void nyalakan(int speed) {
        kecepatan = speed;
        cout << "Motor pin " << pin << " speed: " << kecepatan << endl;
    }
};

int main() {
    MotorDC motorKiri(9);
    motorKiri.nyalakan(150);
    return 0;
}"""
    add_code_block(doc, code_4_1)

    # ==========================================
    # LEVEL 5
    # ==========================================
    add_styled_heading(doc, '🟣 LEVEL 5 — ROBOTIKA & EMBEDDED SYSTEM', 1, RGBColor(128, 0, 128))
    doc.add_paragraph('Penerapan langsung ke hardware. Level ini membahas konsep yang membedakan programmer C++ biasa dengan engineer robotika.')

    add_styled_heading(doc, '5.1 Bit Manipulation (Bahasa Hardware)', 2, RGBColor(128, 0, 128))
    doc.add_paragraph('Di mikrokontroler, kita mengontrol hardware dengan memanipulasi bit individual dalam register. Ini adalah cara tercepat dan paling efisien.')
    
    code_5_1 = """#include <iostream>
using namespace std;

int main() {
    uint8_t reg = 0b00000000; // Register 8-bit, semua mati (0)
    
    // 1. SET bit ke-2 (Nyalakan Pin 2)
    reg |= (1 << 2); 
    
    // 2. CLEAR bit ke-2 (Matikan Pin 2)
    reg &= ~(1 << 2); 
    
    // 3. TOGGLE bit ke-3 (Balik kondisi Pin 3)
    reg ^= (1 << 3); 
    
    // 4. CEK apakah bit ke-3 aktif
    if (reg & (1 << 3)) {
        cout << "Pin 3 sedang AKTIF!" << endl;
    }
    return 0;
}"""
    add_code_block(doc, code_5_1)
    
    add_callout(doc, "Tambahkan keyword 'volatile' pada variabel register hardware (contoh: volatile uint8_t* port). Ini memberi tahu compiler untuk tidak mengoptimasi variabel tersebut.", "warning")

    add_styled_heading(doc, '5.2 Contoh Proyek: Sistem Kontrol Robot', 2, RGBColor(128, 0, 128))
    doc.add_paragraph('Menggabungkan Class, Vector, dan File dalam satu kesatuan.')
    
    code_5_2 = """class Robot {
private:
    string nama;
    float kecepatanMax;
    vector<string> logPerintah;
public:
    Robot(string n, float maxSpeed) : nama(n), kecepatanMax(maxSpeed) {}
    
    void maju(float speed) {
        float s = min(speed, kecepatanMax); // Batasi kecepatan maksimal
        logPerintah.push_back("MAJU speed=" + to_string(s));
    }
    
    void tampilkanLog() {
        for (const auto& log : logPerintah) {
            cout << "[LOG] " << log << endl;
        }
    }
};

int main() {
    Robot r("Robot01", 255.0f);
    r.maju(150.0f);
    r.maju(300.0f); // Akan dibatasi otomatis menjadi 255.0f
    r.tampilkanLog();
    return 0;
}"""
    add_code_block(doc, code_5_2)

    # ==========================================
    # LAMPIRAN
    # ==========================================
    doc.add_page_break()
    add_styled_heading(doc, '🛠️ LAMPIRAN: CHECKLIST & TOOLS', 1, RGBColor(0, 51, 102))
    
    doc.add_paragraph('✅ Checklist Penguasaan Materi (Centang saat sudah paham):')
    checklist = [
        "Level 1: Variabel, I/O, If/Switch, Loop, Fungsi, Array",
        "Level 2: Vector, Map, Algorithm (sort/find), String",
        "Level 3: Pointer, Reference, Dynamic Memory, Smart Pointer",
        "Level 4: Class/Object, Constructor/Destructor, Inheritance, File I/O",
        "Level 5: Bit Manipulation, Multithreading, Serial Comm, Makefile"
    ]
    for item in checklist:
        p = doc.add_paragraph(f"☐  {item}", style='List Bullet')
        p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph('\n🧰 Tools yang Perlu Diinstall (Rekomendasi Ubuntu/Linux):')
    tools_data = [
        ["build-essential", "Compiler g++ dan make", "sudo apt install build-essential"],
        ["git", "Version control", "sudo apt install git"],
        ["VS Code", "Code Editor terbaik", "sudo snap install code --classic"],
        ["Arduino IDE", "Programming board Arduino", "sudo snap install arduino"],
        ["ROS2 Humble", "Robot Operating System", "Lihat docs.ros.org"]
    ]
    add_styled_table(doc, ["Tool", "Kegunaan", "Perintah Install"], tools_data)

    doc.add_paragraph('\n📚 Sumber Belajar Lanjutan:')
    sources = [
        "learncpp.com → Tutorial C++ dari nol paling lengkap dan gratis (Sangat Direkomendasikan).",
        "cppreference.com → Kamus referensi resmi semua fitur C++.",
        "docs.arduino.cc → Dokumentasi resmi fungsi-fungsi Arduino.",
        "docs.ros.org → Dokumentasi resmi ROS2."
    ]
    for src in sources:
        doc.add_paragraph(src, style='List Bullet')

    # Simpan file
    filename = "Modul_CPP_Robotika_Sempurna.docx"
    doc.save(filename)
    print(f"✅ Sukses! File '{filename}' berhasil dibuat dengan format yang rapi dan profesional.")

if __name__ == "__main__":
    main()