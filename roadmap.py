from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_color(doc, text, level, color_rgb):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color_rgb
    return heading

def create_roadmap():
    doc = Document()
    
    # --- STYLE SETUP ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- TITLE ---
    title = doc.add_heading('🚀 ROADMAP HEKA: MENUJU ELEKTRO UM 2026', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102) # Navy Blue

    subtitle = doc.add_paragraph('"Building The Future, One Line of Code at a Time."')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    # --- SECTION 1: VISI ---
    add_heading_with_color(doc, '🎯 VISI & MISI', 1, RGBColor(0, 51, 102))
    p = doc.add_paragraph()
    p.add_run('Target Utama: ').bold = True
    p.add_run('Lolos Teknik Elektro Universitas Negeri Malang (2026)\n')
    p.add_run('Identitas: ').bold = True
    p.add_run('Programmer • Engineer • Content Creator')

    # --- SECTION 2: KESEHATAN ---
    add_heading_with_color(doc, '🛡️ PILAR 1: KESEHATAN & MENTAL', 1, RGBColor(34, 139, 34)) # Forest Green
    doc.add_paragraph('😴 Tidur maksimal 22.30 (Wajib)')
    doc.add_paragraph('💧 Air putih 2–3 liter/hari')
    doc.add_paragraph('🚶 Bergerak 5 menit setiap 1 jam duduk')
    doc.add_paragraph('🧘 Kurangi overthinking & Shorts')

    # --- SECTION 3: AKADEMIK ---
    add_heading_with_color(doc, '🎓 PILAR 2: AKADEMIK (TIKET MASUK UM)', 1, RGBColor(0, 0, 139)) # Dark Blue
    
    p = doc.add_paragraph()
    p.add_run('🥇 MATEMATIKA').bold = True
    doc.add_paragraph('Aljabar, Fungsi, Trigonometri, Vektor, Kalkulus Dasar.', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('🥈 FISIKA').bold = True
    doc.add_paragraph('Listrik & Elektromagnetik (Prioritas Utama), Gerak, Energi.', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('🥉 KIMIA').bold = True
    doc.add_paragraph('Fokus lulus dan paham dasar material.', style='List Bullet')

    # --- SECTION 4: HARD SKILLS ---
    add_heading_with_color(doc, '⚡ PILAR 3: HARD SKILLS', 1, RGBColor(204, 0, 0)) # Red
    
    # C++
    p = doc.add_paragraph()
    p.add_run('💻 C++ (Algoritma & Lomba)').bold = True
    doc.add_paragraph('Level 1-7: Variabel s.d. Dynamic Programming & Dijkstra.', style='List Bullet')
    doc.add_paragraph('Target: TLX & Lomba Informatika.', style='List Bullet')

    # Python
    p = doc.add_paragraph()
    p.add_run('🐍 Python (Proyek & IoT)').bold = True
    doc.add_paragraph('Skill: OOP, API, Automation.', style='List Bullet')
    doc.add_paragraph('Proyek: Habit Tracker, Discord Bot, Dashboard ESP32.', style='List Bullet')

    # Web
    p = doc.add_paragraph()
    p.add_run('🌐 Web Dev (Portofolio)').bold = True
    doc.add_paragraph('Stack: HTML, CSS, JS (DOM/Async).', style='List Bullet')

    # Robotika
    p = doc.add_paragraph()
    p.add_run('🤖 Robotika (Elektro)').bold = True
    doc.add_paragraph('Hardware: Arduino ➡️ ESP32.', style='List Bullet')
    doc.add_paragraph('Proyek: Smart Room, Smart Garden.', style='List Bullet')

    # --- SECTION 5: ROUTINE TABLE ---
    add_heading_with_color(doc, '⏳ PROTOKOL HARIAN', 1, RGBColor(128, 0, 128)) # Purple
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Aktivitas'
    hdr_cells[1].text = 'Durasi'
    hdr_cells[2].text = 'Fokus'
    
    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    data = [
        ('Akademik', '60 Menit', 'Matematika / Fisika'),
        ('Coding (C++)', '30 Menit', 'Latihan Algoritma'),
        ('Dev (Py/Web)', '30 Menit', 'Proyek / Skill Baru'),
        ('Robotika', '15 Menit', 'Eksperimen Alat'),
        ('Reading', '10 Menit', 'Buku/Artikel Tech'),
        ('Rest', '-', 'Tidur < 22.30')
    ]

    for act, dur, foc in data:
        row_cells = table.add_row().cells
        row_cells[0].text = act
        row_cells[1].text = dur
        row_cells[2].text = foc

    # --- FOOTER ---
    doc.add_page_break()
    motto = doc.add_paragraph('"Jangan belajar hanya untuk sekolah. Belajarlah untuk membangun masa depan."')
    motto.alignment = WD_ALIGN_PARAGRAPH.CENTER
    motto.runs[0].font.size = Pt(14)
    motto.runs[0].italic = True
    motto.runs[0].font.color.rgb = RGBColor(255, 102, 0)

    doc.save('Roadmap_HEKA_2026.docx')
    print("✅ File 'Roadmap_HEKA_2026.docx' berhasil dibuat!")

if __name__ == "__main__":
    create_roadmap()